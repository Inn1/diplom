from flask import Blueprint, render_template, request, redirect, url_for, flash, current_app, session, abort, send_from_directory, jsonify
from models import db, User, Role, Category, Document, DocumentHistory  
from flask_login import login_user, logout_user, login_required, current_user
from PIL import Image
import pytesseract
import os
import io
import fitz 
import re
from rapidfuzz import fuzz
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
from docx import Document as DocxDocument
from werkzeug.utils import secure_filename
import re, hashlib, string 
import unicodedata
import logging
from sqlalchemy.orm import joinedload
from sqlalchemy import func, or_
from datetime import datetime
logging.basicConfig(level=logging.DEBUG)


main = Blueprint('main', __name__)


@main.before_request
def require_login():
    # Разрешаем доступ только к странице входа 
    allowed_routes = ['main.login']  #имена маршрутов
    if request.endpoint not in allowed_routes and not current_user.is_authenticated:
        return redirect(url_for('main.login', next=request.url))
    
@main.route('/documents/<int:document_id>')
def document_detail(document_id):
    doc = Document.query.get_or_404(document_id)
    return render_template('document_detail.html', document=doc)

@main.route('/documents')
def documents_list():
    page = request.args.get('page', 1, type=int)  # Получаем номер текущей страницы из параметров URL, по умолчанию 1
    per_page = 10  # Количество документов на странице, можно изменить под ваши нужды

    pagination = Document.query.order_by(Document.date_added.desc()).paginate(page=page, per_page=per_page)
    documents = pagination.items  # Документы на текущей странице
    total_pages = pagination.pages  # Общее количество страниц

    return render_template('index.html',
                           documents=documents,
                           total_pages=total_pages,
                           page=page)


@main.route('/index')
@main.route('/')
def index():
    q = request.args.get('q', '', type=str).strip()
    category_id = request.args.get('category_id', type=int)
    date_from = request.args.get('date_from', type=str)
    date_to = request.args.get('date_to', type=str)

    query = Document.query.options(
        joinedload(Document.user),
        joinedload(Document.category)
    )

    # Заменяем множественные пробелы одним пробелом
    q = re.sub(r'\s+', ' ', q)
    # Разбиваем строку по несловесным символам на отдельные слова
    search_terms = re.split(r'\W+', q)

    for term in search_terms:
        if term:
            pattern = f"%{term}%"
            query = query.filter(
            or_(
            Document.filename.ilike(pattern),
            Document.text.ilike(pattern)
            )
        )
    if category_id:
        query = query.filter(Document.category_id == category_id)

    if date_from:
        try:
            dt_from = datetime.strptime(date_from, '%Y-%m-%d').date()
            query = query.filter(Document.date_added >= dt_from)
        except ValueError:
            pass

    if date_to:
        try:
            dt_to = datetime.strptime(date_to, '%Y-%m-%d').date()
            query = query.filter(Document.date_added <= dt_to)
        except ValueError:
            pass

    page = request.args.get('page', 1, type=int)
    per_page = 10  # сколько документов на странице показывать

    # Пагинация применяется к отфильтрованному запросу
    pagination = query.order_by(Document.date_added.desc()).paginate(page=page, per_page=per_page)
    documents = pagination.items
    total_pages = pagination.pages

    categories = Category.query.order_by(Category.name).all()

    return render_template('index.html', documents=documents, categories=categories, page=page, total_pages=total_pages)

@main.route('/uploads/<filename>')
def uploaded_file(filename):
    uploads = current_app.config['UPLOAD_FOLDER']
    file_path = os.path.join(uploads, filename)
    if not os.path.exists(file_path):
        abort(404)
    return send_from_directory(uploads, filename)


@main.route('/profile')
def profile():
   # Получаем документы, загруженные текущим пользователем
    user_documents = Document.query.filter_by(user_id=current_user.id).options(
        joinedload(Document.user),
        joinedload(Document.category)
    ).all()

    
    page = request.args.get('page', 1, type=int)
    per_page = 10  # сколько документов на странице показывать

    pagination = Document.query.order_by(Document.date_added.desc()).paginate(page=page, per_page=per_page)
    documents = pagination.items
    total_pages = pagination.pages

    return render_template('profile.html', documents=user_documents, page=page, total_pages=total_pages)

#ДОБАВИТЬ ДОКУМЕНТ

def pdf_to_text_pymupdf(pdf_bytes):
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    texts = []
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        pix = page.get_pixmap()
        img_data = pix.tobytes("png")
        image = Image.open(io.BytesIO(img_data))
        text = pytesseract.image_to_string(image, lang='rus+eng')
        texts.append(text)
    return "\n\n--- Страница ---\n\n".join(texts)


def extract_text_from_pdf(pdf_bytes):
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    all_text = ""
    for page in doc:
        all_text += page.get_text() + "\n"
    return all_text

#Имена файлов
def secure_filename_ru(filename):
    filename = unicodedata.normalize('NFKC', filename)
    # Разрешаем буквы (включая русские), цифры, дефисы, подчёркивания и точки
    filename = re.sub(r'[^a-zA-Z0-9а-яА-ЯёЁ\-_\.]', '_', filename)
    return filename


def get_unique_filename(directory, filename):
    """
    Проверяет, существует ли файл с именем filename в директории.
    Если существует, добавляет суффикс _1, _2 и т.д., чтобы имя было уникальным.
    """
    base, ext = os.path.splitext(filename)
    counter = 1
    unique_name = filename
    while os.path.exists(os.path.join(directory, unique_name)):
        unique_name = f"{base}_{counter}{ext}"
        counter += 1
    return unique_name



@main.route('/addDoc', methods=['GET', 'POST'])
def addDoc():
    if request.method == 'POST':
        file = request.files.get('file')

        if not file:
            flash('Пожалуйста, выберите файл.', 'warning')
            return render_template('documents/addDoc.html')

        filename_secure = secure_filename_ru(file.filename)
        temp_path = os.path.join(current_app.config['TEMP_UPLOAD_FOLDER'], filename_secure)

        file.save(temp_path)

        extracted_text = ''
        try:
            if filename_secure.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.tiff')):
                image = Image.open(temp_path)
                extracted_text = pytesseract.image_to_string(image, lang='rus+eng')

            elif filename_secure.lower().endswith('.pdf'):
                with open(temp_path, 'rb') as f:
                    pdf_bytes = f.read()
                doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                for page in doc:
                    extracted_text += page.get_text() + '\n'

            elif filename_secure.lower().endswith('.txt'):
                with open(temp_path, 'r', encoding='utf-8') as f:
                    extracted_text = f.read()

            elif filename_secure.lower().endswith('.docx'):
                doc = DocxDocument(temp_path)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                extracted_text = '\n'.join(full_text)

            else:
                flash('Недопустимый формат файла.', 'danger')
                os.remove(temp_path)
                return render_template('documents/addDoc.html')

            # Сохраняем распознанный текст в файл
            text_filename = filename_secure + '.txt'
            text_file_path = os.path.join(current_app.config['TEMP_UPLOAD_FOLDER'], text_filename)
            with open(text_file_path, 'w', encoding='utf-8') as f:
                f.write(extracted_text)

            # Сохраняем пути к файлам в сессию
            session['temp_file_path'] = temp_path
            session['temp_text_path'] = text_file_path

        except Exception as e:
            flash(f'Ошибка при обработке файла: {e}', 'danger')
            os.remove(temp_path)
            return render_template('documents/addDoc.html')

        print("Saved to session:", session['temp_file_path'], session['temp_text_path'])
        return redirect(url_for('main.confirmDoc'))

    return render_template('documents/addDoc.html')


def normalize_text(text):
    text = text.lower()
    text = text.translate(str.maketrans('', '', string.punctuation))
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def compute_text_hash(text):
    norm = normalize_text(text)
    return hashlib.sha256(norm.encode('utf-8')).hexdigest()



def is_duplicate(text1, text2, threshold=85, max_diff_chars=300):
    # Приводим к нижнему регистру и убираем лишние пробелы
    clean1 = text1.lower().strip()
    clean2 = text2.lower().strip()

    # token_set_ratio() для нечувствительности к порядку слов
    # ratio() для точного сравнения
    ratio_set = fuzz.token_set_ratio(clean1, clean2)
    ratio_exact = fuzz.ratio(clean1, clean2)

    # Определяем длину максимальной разницы
    #diff_len = abs(len(clean1) - len(clean2))
    # Проверка основного условия
    if ratio_set >= threshold and ratio_exact >= threshold:
        #if diff_len > max_diff_chars:
            #return False
        return True
    return False


@main.route('/addDoc/confirm', methods=['GET', 'POST'])
@login_required
def confirmDoc():
    temp_file_path = session.get('temp_file_path')
    temp_text_path = session.get('temp_text_path')
    categories = Category.query.all()

    if not temp_file_path or not os.path.exists(temp_file_path):
        flash('Временный файл не найден. Пожалуйста, загрузите файл заново.', 'warning')
        return redirect(url_for('main.addDoc'))

    if not temp_text_path or not os.path.exists(temp_text_path):
        flash('Распознанный текст не найден. Пожалуйста, загрузите файл заново.', 'warning')
        return redirect(url_for('main.addDoc'))

    with open(temp_text_path, 'r', encoding='utf-8') as f:
        extracted_text = f.read()

    if request.method == 'POST':
        edited_text = request.form.get('text')
        category_id = request.form.get('category')
        force_add = request.form.get('force_add')  # Новый параметр для подтверждения добавления

        if not edited_text or not category_id:
            flash('Пожалуйста, заполните все поля.', 'warning')
            return render_template('documents/confirmDoc.html', text=edited_text, categories=categories, selected_category=category_id)

        # Поиск похожих документов
        similar_docs = []
        documents = Document.query.all()
        for doc in documents:
            if is_duplicate(edited_text, doc.text):
                similar_docs.append({'filename': doc.filename, 'id': doc.id})

        if similar_docs and not force_add:
            # Показываем предупреждение с похожими документами
            return render_template('documents/confirmDoc.html',
                                   text=edited_text,
                                   categories=categories,
                                   selected_category=category_id,
                                   similar_docs=similar_docs,
                                   warning='Найдены похожие документы. Вы уверены, что хотите добавить?',
                                   show_modal=True)

        # Если пользователь подтвердил добавление (force_add) — сохраняем документ
        upload_folder = current_app.config['UPLOAD_FOLDER']
        os.makedirs(upload_folder, exist_ok=True)
        filename_unique = get_unique_filename(upload_folder, os.path.basename(temp_file_path))
        permanent_path = os.path.join(upload_folder, filename_unique)
        os.rename(temp_file_path, permanent_path)

        doc = Document(
            filename=filename_unique,
            filepath=permanent_path,
            text=edited_text,
            category_id=int(category_id),
            user_id=current_user.user_id
        )
        db.session.add(doc)
        db.session.commit()

        if os.path.exists(temp_text_path):
            os.remove(temp_text_path)

        session.pop('temp_file_path', None)
        session.pop('temp_text_path', None)

        flash('Документ успешно сохранён!', 'success')
        return redirect(url_for('main.addDoc'))

    # GET-запрос — просто отдаем форму с распознанным текстом
    return render_template('documents/confirmDoc.html', text=extracted_text, categories=categories, selected_category=None)



@main.route('/documents/delete/<int:document_id>', methods=['POST'])
@login_required
def delete_document(document_id):
    doc = Document.query.get_or_404(document_id)

    # Проверяем права доступа
    if not (current_user.role.name == 'admin' or current_user.user_id == doc.user_id):
        abort(403)  # Запрещено
    
    # Удаляем связанные записи истории
    DocumentHistory.query.filter_by(document_id=document_id).delete()

    # Удаляем файл с диска, если нужно
    if os.path.exists(doc.filepath):
        os.remove(doc.filepath)

    # Удаляем запись из БД
    db.session.delete(doc)
    db.session.commit()

    flash('Документ успешно удалён.', 'success')
    return redirect(url_for('main.documents_list'))



@main.route('/documents/edit/<int:document_id>', methods=['GET', 'POST'])
def edit_document(document_id):
    if current_user.role.name != 'admin':
        abort(403)  # Доступ только администратору

    doc = Document.query.get_or_404(document_id)

    if request.method == 'POST':
        filename = request.form.get('filename', '').strip()
        text = request.form.get('text', '').strip()
        category_id = request.form.get('category_id')

        # Валидация
        if not filename:
            flash('Имя файла не может быть пустым', 'danger')
            categories = Category.query.all()
            return render_template('edit_document.html', doc=doc, categories=categories)

        # --- Добавленная проверка на пустой текст ---
        if not text:
            flash('Текст документа не может быть пустым', 'warning')
            categories = Category.query.all()
            return render_template('edit_document.html', doc=doc, categories=categories)
        # --- Конец добавленной проверки ---

        # --- Добавлено: сохраняем историю изменений, если текст изменился ---
        if doc.text != text:
            history_entry = DocumentHistory(
                document_id=doc.id,
                user_id=current_user.user_id,
                old_text=doc.text,
                new_text=text,
                changed_at=datetime.utcnow()
            )
            db.session.add(history_entry)
        # --- Конец добавления истории изменений ---

        # Обновляем поля документа
        doc.filename = filename
        doc.text = text
        if category_id:
            doc.category_id = int(category_id)
        else:
            doc.category_id = None

        db.session.commit()

        flash('Документ успешно обновлён', 'success')
        return redirect(url_for('main.document_detail', document_id=doc.id))

    # GET - показываем форму с текущими данными
    categories = Category.query.all()
    return render_template('edit_document.html', doc=doc, categories=categories)

#маршрут для просмотра истории изменений документа
@main.route('/documents/<int:document_id>/history')
def document_history(document_id):
    # Получаем документ или 404
    doc = Document.query.get_or_404(document_id)
    
    # Получаем все записи истории для документа, сортируем по дате (новые сверху)
    history = DocumentHistory.query.filter_by(document_id=document_id).order_by(DocumentHistory.changed_at.desc()).all()
    
    return render_template('document_history.html', document=doc, history=history)

@main.route('/documents/<int:document_id>/history/<int:history_id>')
def document_history_detail(document_id, history_id):
    doc = Document.query.get_or_404(document_id)
    history_entry = DocumentHistory.query.get_or_404(history_id)

    if history_entry.document_id != document_id:
        abort(404)  # Проверка, что запись истории относится к этому документу

    return render_template('document_history_detail.html', document=doc, history_entry=history_entry)


@main.route('/login', methods = ['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('main.index'))
    
    error = None  # переменная для ошибки
    
    if request.method == 'POST':
        login_input = request.form['login']
        password_input = request.form['password']

        user = User.query.filter_by(login=login_input).first()
        if user and user.check_password(password_input):
            login_user(user)
            next_page = request.args.get('next')
            return redirect(next_page or url_for('main.index'))
        else:
            error = 'Неверный логин или пароль'  # записываем ошибку

    return render_template('login.html', error=error)


@main.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('main.login'))



@main.route('/about')
def about():
    return render_template('about.html')
